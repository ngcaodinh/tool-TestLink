"""
TestLink Report Generator - Creates Word documents from test plan data
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List, Any
from datetime import datetime


class ReportGenerator:
    def __init__(self, client):
        self.client = client
        self.doc = None
        
    def generate_report(self, plan_id: int, build_id: int = None, output_path: str = None) -> str:
        self.doc = Document()
        self._setup_styles()
        
        plan_info = self._get_plan_info(plan_id)
        if not plan_info:
            raise ValueError(f"Test plan with ID {plan_id} not found")
        
        self._add_title(f"Test Execution Report: {plan_info['name']}")
        self._add_meta_info(plan_info, build_id)
        
        metrics = self.client.get_plan_metrics(plan_id, build_id)
        testcases = self.client.get_testcases_in_plan(plan_id, build_id)
        suite_data = self._get_suite_structure(plan_id)
        
        self._add_table_of_contents()
        self._add_suite_descriptions(suite_data)
        self._add_testcase_summary(testcases)
        self._add_testcase_scenarios(testcases)
        self._add_testcase_details(testcases)
        self._add_testcase_keywords(testcases)
        self._add_testcase_custom_fields(testcases)
        self._add_requirements_coverage(testcases)
        self._add_execution_notes(plan_id, build_id)
        self._add_step_execution_notes(plan_id, build_id)
        self._add_test_results(testcases)
        self._add_step_results(plan_id, build_id)
        self._add_build_custom_fields(build_id)
        self._add_metrics(metrics, testcases)
        
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"test_report_{plan_id}_{timestamp}.docx"
        
        self.doc.save(output_path)
        return output_path
    
    def _setup_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        for i in range(1, 4):
            try:
                heading_style = self.doc.styles[f'Heading {i}']
                heading_style.font.name = 'Calibri'
                heading_style.font.size = Pt(16 - (i * 2))
                heading_style.font.bold = True
            except:
                pass
    
    def _add_title(self, title: str):
        heading = self.doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_meta_info(self, plan_info: Dict, build_id: int):
        self.doc.add_heading('Report Information', level=1)
        
        table = self.doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        table.rows[0].cells[0].text = 'Test Plan'
        table.rows[0].cells[1].text = plan_info.get('name', 'N/A')
        
        table.rows[1].cells[0].text = 'Project'
        table.rows[1].cells[1].text = plan_info.get('project_name', 'N/A')
        
        table.rows[2].cells[0].text = 'Generated'
        table.rows[2].cells[1].text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        table.rows[3].cells[0].text = 'Build'
        if build_id:
            build_info = self._get_build_info(build_id)
            table.rows[3].cells[1].text = build_info.get('name', f'Build #{build_id}')
        else:
            table.rows[3].cells[1].text = 'All Builds'
        
        self.doc.add_paragraph()
    
    def _add_table_of_contents(self):
        self.doc.add_heading('Table of Contents', level=1)
        
        toc_items = [
            '1. Test Suite Descriptions',
            '2. Test Case Summary',
            '3. Test Case Scenarios',
            '4. Test Case Authors',
            '5. Test Case Keywords',
            '6. Test Case Custom Fields',
            '7. Test Case Related Requirements',
            '8. Execution Notes',
            '9. Step Execution Notes',
            '10. Test Results',
            '11. Step Execution Results',
            '12. Build Custom Fields',
            '13. Metrics Summary'
        ]
        
        for item in toc_items:
            p = self.doc.add_paragraph(item)
            p.paragraph_format.space_after = Pt(2)
    
    def _add_suite_descriptions(self, suite_data: List[Dict]):
        self.doc.add_heading('Test Suite Descriptions', level=1)
        
        if not suite_data:
            self.doc.add_paragraph('No test suites found.')
            return
        
        for suite in suite_data:
            self.doc.add_heading(suite.get('name', 'Unnamed Suite'), level=2)
            
            details = suite.get('details', '')
            if details:
                self.doc.add_paragraph(details)
            else:
                self.doc.add_paragraph('No description available.')
    
    def _add_testcase_summary(self, testcases: List[Dict]):
        self.doc.add_heading('Test Case Summary', level=1)
        
        if not testcases:
            self.doc.add_paragraph('No test cases found.')
            return
        
        table = self.doc.add_table(rows=len(testcases) + 1, cols=5)
        table.style = 'Table Grid'
        
        headers = ['External ID', 'Test Case Name', 'Version', 'Summary', 'Status']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self._set_cell_background(cell, '4472C4')
            self._set_cell_text_color(cell, 'FFFFFF')
        
        for row_idx, tc in enumerate(testcases, 1):
            table.rows[row_idx].cells[0].text = str(tc.get('tc_id', ''))
            table.rows[row_idx].cells[1].text = tc.get('name', 'N/A')
            table.rows[row_idx].cells[2].text = str(tc.get('version', ''))
            summary = tc.get('summary', '')
            if summary and len(summary) > 50:
                summary = summary[:50] + '...'
            table.rows[row_idx].cells[3].text = summary or 'N/A'
            table.rows[row_idx].cells[4].text = self._format_status(tc.get('status'))
        
        self.doc.add_paragraph()
    
    def _add_testcase_scenarios(self, testcases: List[Dict]):
        self.doc.add_heading('Test Case Scenarios', level=1)
        
        if not testcases:
            self.doc.add_paragraph('No test cases found.')
            return
        
        for tc in testcases:
            self.doc.add_heading(tc.get('name', 'Unnamed TC'), level=2)
            
            steps = self.client.get_tc_steps(tc.get('version_id'))
            
            if steps:
                for step in steps:
                    p = self.doc.add_paragraph()
                    p.style = 'List Number'
                    p.add_run(f"Step {step['step_number']}: ").bold = True
                    p.add_run(step.get('actions', ''))
                    
                    if step.get('expected_results'):
                        p2 = self.doc.add_paragraph()
                        p2.paragraph_format.left_indent = Inches(0.5)
                        p2.add_run("Expected: ").italic = True
                        p2.add_run(step.get('expected_results', ''))
            else:
                self.doc.add_paragraph('No steps defined.')
    
    def _add_testcase_details(self, testcases: List[Dict]):
        self.doc.add_heading('Test Case Authors', level=1)
        
        if not testcases:
            self.doc.add_paragraph('No test cases found.')
            return
        
        for tc in testcases:
            self.doc.add_heading(tc.get('name', 'Unnamed TC'), level=2)
            
            author = self.client.get_tc_author(tc.get('tc_id'))
            if author:
                self.doc.add_paragraph(f"Author: {author.get('login', 'Unknown')}")
                name_parts = []
                if author.get('first_name'):
                    name_parts.append(author['first_name'])
                if author.get('last_name'):
                    name_parts.append(author['last_name'])
                if name_parts:
                    self.doc.add_paragraph(f"Name: {' '.join(name_parts)}")
                if author.get('email'):
                    self.doc.add_paragraph(f"Email: {author['email']}")
            else:
                self.doc.add_paragraph('Author information not available.')
    
    def _add_testcase_keywords(self, testcases: List[Dict]):
        self.doc.add_heading('Test Case Keywords', level=1)
        
        if not testcases:
            self.doc.add_paragraph('No test cases found.')
            return
        
        for tc in testcases:
            self.doc.add_heading(tc.get('name', 'Unnamed TC'), level=2)
            
            keywords = self.client.get_tc_keywords(tc.get('tc_id'))
            
            if keywords:
                for kw in keywords:
                    self.doc.add_paragraph(f"• {kw.get('keyword', '')}", style='List Bullet')
            else:
                self.doc.add_paragraph('No keywords assigned.')
    
    def _add_testcase_custom_fields(self, testcases: List[Dict]):
        self.doc.add_heading('Test Case Custom Fields', level=1)
        
        if not testcases:
            self.doc.add_paragraph('No test cases found.')
            return
        
        for tc in testcases:
            self.doc.add_heading(tc.get('name', 'Unnamed TC'), level=2)
            
            custom_fields = self.client.get_tc_custom_fields(tc.get('tc_id'), tc.get('version_id'))
            
            if custom_fields:
                table = self.doc.add_table(rows=len(custom_fields) + 1, cols=2)
                table.style = 'Table Grid'
                table.rows[0].cells[0].text = 'Field Name'
                table.rows[0].cells[1].text = 'Value'
                self._set_cell_background(table.rows[0].cells[0], '4472C4')
                self._set_cell_background(table.rows[0].cells[1], '4472C4')
                self._set_cell_text_color(table.rows[0].cells[0], 'FFFFFF')
                self._set_cell_text_color(table.rows[0].cells[1], 'FFFFFF')
                
                for idx, cf in enumerate(custom_fields, 1):
                    table.rows[idx].cells[0].text = cf.get('label', cf.get('name', ''))
                    table.rows[idx].cells[1].text = cf.get('value', 'N/A')
            else:
                self.doc.add_paragraph('No custom fields defined.')
    
    def _add_requirements_coverage(self, testcases: List[Dict]):
        self.doc.add_heading('Test Case Related Requirements', level=1)
        
        if not testcases:
            self.doc.add_paragraph('No test cases found.')
            return
        
        for tc in testcases:
            self.doc.add_heading(tc.get('name', 'Unnamed TC'), level=2)
            
            requirements = self.client.get_requirements_coverage(tc.get('tc_id'))
            
            if requirements:
                table = self.doc.add_table(rows=len(requirements) + 1, cols=3)
                table.style = 'Table Grid'
                headers = ['Req ID', 'Title', 'Spec']
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h
                    self._set_cell_background(table.rows[0].cells[i], '4472C4')
                    self._set_cell_text_color(table.rows[0].cells[i], 'FFFFFF')
                
                for idx, req in enumerate(requirements, 1):
                    table.rows[idx].cells[0].text = req.get('req_doc_id', '')
                    table.rows[idx].cells[1].text = req.get('title', '')
                    table.rows[idx].cells[2].text = req.get('spec_title', '')
            else:
                self.doc.add_paragraph('No requirements mapped.')
    
    def _add_execution_notes(self, plan_id: int, build_id: int = None):
        self.doc.add_heading('Execution Notes', level=1)
        
        notes = self.client.get_execution_notes(plan_id, build_id)
        
        if not notes:
            self.doc.add_paragraph('No execution notes found.')
            return
        
        table = self.doc.add_table(rows=len(notes) + 1, cols=5)
        table.style = 'Table Grid'
        headers = ['Test Case', 'Build', 'Tester', 'Status', 'Notes']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            self._set_cell_background(table.rows[0].cells[i], '4472C4')
            self._set_cell_text_color(table.rows[0].cells[i], 'FFFFFF')
        
        for idx, note in enumerate(notes, 1):
            table.rows[idx].cells[0].text = note.get('tc_name', '')
            table.rows[idx].cells[1].text = note.get('build_name', '')
            table.rows[idx].cells[2].text = note.get('tester', '')
            table.rows[idx].cells[3].text = self._format_status(note.get('status'))
            notes_text = note.get('notes', '') or ''
            if len(notes_text) > 100:
                notes_text = notes_text[:100] + '...'
            table.rows[idx].cells[4].text = notes_text
        
        self.doc.add_paragraph()
    
    def _add_step_execution_notes(self, plan_id: int, build_id: int = None):
        self.doc.add_heading('Step Execution Notes', level=1)
        
        executions = self.client.get_all_plan_executions(plan_id, build_id)
        
        if not executions:
            self.doc.add_paragraph('No step execution notes found.')
            return
        
        for exec_data in executions[:20]:
            self.doc.add_heading(exec_data.get('tc_name', 'Unknown TC'), level=2)
            
            steps = self.client.get_step_results(exec_data.get('id'))
            
            if steps:
                table = self.doc.add_table(rows=len(steps) + 1, cols=4)
                table.style = 'Table Grid'
                headers = ['Step', 'Action', 'Status', 'Notes']
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h
                    self._set_cell_background(table.rows[0].cells[i], '4472C4')
                    self._set_cell_text_color(table.rows[0].cells[i], 'FFFFFF')
                
                for idx, step in enumerate(steps, 1):
                    table.rows[idx].cells[0].text = str(step.get('step_number', ''))
                    action = step.get('actions', '')
                    if len(action) > 50:
                        action = action[:50] + '...'
                    table.rows[idx].cells[1].text = action
                    table.rows[idx].cells[2].text = self._format_status(step.get('status'))
                    table.rows[idx].cells[3].text = step.get('notes', '') or ''
            else:
                self.doc.add_paragraph('No step results recorded.')
    
    def _add_test_results(self, testcases: List[Dict]):
        self.doc.add_heading('Test Results', level=1)
        
        if not testcases:
            self.doc.add_paragraph('No test results found.')
            return
        
        table = self.doc.add_table(rows=len(testcases) + 1, cols=4)
        table.style = 'Table Grid'
        headers = ['Test Case', 'Version', 'Status', 'Tester']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            self._set_cell_background(table.rows[0].cells[i], '4472C4')
            self._set_cell_text_color(table.rows[0].cells[i], 'FFFFFF')
        
        for idx, tc in enumerate(testcases, 1):
            status = tc.get('status')
            table.rows[idx].cells[0].text = tc.get('name', '')
            table.rows[idx].cells[1].text = str(tc.get('version', ''))
            table.rows[idx].cells[2].text = self._format_status(status)
            table.rows[idx].cells[3].text = tc.get('tester', '')
            
            if status:
                status_color = self._get_status_color(status)
                self._set_cell_background(table.rows[idx].cells[2], status_color)
                self._set_cell_text_color(table.rows[idx].cells[2], 'FFFFFF')
        
        self.doc.add_paragraph()
    
    def _add_step_results(self, plan_id: int, build_id: int = None):
        self.doc.add_heading('Step Execution Results', level=1)
        
        executions = self.client.get_all_plan_executions(plan_id, build_id)
        
        if not executions:
            self.doc.add_paragraph('No step execution results found.')
            return
        
        has_results = False
        for exec_data in executions:
            steps = self.client.get_step_results(exec_data.get('id'))
            if steps:
                has_results = True
                break
        
        if not has_results:
            self.doc.add_paragraph('No step-level results recorded.')
            return
        
        for exec_data in executions:
            steps = self.client.get_step_results(exec_data.get('id'))
            if not steps:
                continue
                
            self.doc.add_heading(exec_data.get('tc_name', 'Unknown TC'), level=2)
            
            table = self.doc.add_table(rows=len(steps) + 1, cols=3)
            table.style = 'Table Grid'
            headers = ['Step', 'Status', 'Notes']
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                self._set_cell_background(table.rows[0].cells[i], '4472C4')
                self._set_cell_text_color(table.rows[0].cells[i], 'FFFFFF')
            
            for idx, step in enumerate(steps, 1):
                table.rows[idx].cells[0].text = str(step.get('step_number', ''))
                table.rows[idx].cells[1].text = self._format_status(step.get('status'))
                table.rows[idx].cells[2].text = step.get('notes', '') or ''
                
                status = step.get('status')
                if status:
                    status_color = self._get_status_color(status)
                    self._set_cell_background(table.rows[idx].cells[1], status_color)
                    self._set_cell_text_color(table.rows[idx].cells[1], 'FFFFFF')
    
    def _add_build_custom_fields(self, build_id: int = None):
        self.doc.add_heading('Build Custom Fields', level=1)
        
        if not build_id:
            self.doc.add_paragraph('No specific build selected.')
            return
        
        custom_fields = self.client.get_build_custom_fields(build_id)
        
        if not custom_fields:
            self.doc.add_paragraph('No custom fields for this build.')
            return
        
        table = self.doc.add_table(rows=len(custom_fields) + 1, cols=2)
        table.style = 'Table Grid'
        headers = ['Field Name', 'Value']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            self._set_cell_background(table.rows[0].cells[i], '4472C4')
            self._set_cell_text_color(table.rows[0].cells[i], 'FFFFFF')
        
        for idx, cf in enumerate(custom_fields, 1):
            table.rows[idx].cells[0].text = cf.get('label', cf.get('name', ''))
            table.rows[idx].cells[1].text = cf.get('value', 'N/A')
    
    def _add_metrics(self, metrics: Dict, testcases: List[Dict]):
        self.doc.add_heading('Metrics Summary', level=1)
        
        total_tcs = len(testcases)
        executed_tcs = len([tc for tc in testcases if tc.get('status')])
        
        metrics_table = self.doc.add_table(rows=8, cols=2)
        metrics_table.style = 'Table Grid'
        
        data = [
            ('Total Test Cases', str(metrics.get('total_tcs', total_tcs))),
            ('Executed Test Cases', str(metrics.get('total_tcs', executed_tcs))),
            ('Passed', str(metrics.get('passed', 0))),
            ('Failed', str(metrics.get('failed', 0))),
            ('Blocked', str(metrics.get('blocked', 0))),
            ('Warning', str(metrics.get('warning', 0))),
            ('Not Run', str(metrics.get('not_run', 0))),
            ('Total Builds', str(metrics.get('total_builds', 0)))
        ]
        
        for idx, (label, value) in enumerate(data):
            metrics_table.rows[idx].cells[0].text = label
            metrics_table.rows[idx].cells[1].text = value
        
        self.doc.add_paragraph()
        
        if total_tcs > 0:
            pass_rate = (metrics.get('passed', 0) / total_tcs) * 100
            fail_rate = (metrics.get('failed', 0) / total_tcs) * 100
            blocked_rate = (metrics.get('blocked', 0) / total_tcs) * 100
            
            rate_table = self.doc.add_table(rows=3, cols=2)
            rate_table.style = 'Table Grid'
            
            rate_data = [
                ('Pass Rate', f'{pass_rate:.1f}%'),
                ('Fail Rate', f'{fail_rate:.1f}%'),
                ('Blocked Rate', f'{blocked_rate:.1f}%')
            ]
            
            for idx, (label, value) in enumerate(rate_data):
                rate_table.rows[idx].cells[0].text = label
                rate_table.rows[idx].cells[1].text = value
    
    def _get_plan_info(self, plan_id: int) -> Dict:
        query = """
            SELECT tp.id, nh_plan.name, tp.notes, nh_proj.name as project_name, tp.testproject_id
            FROM testplans tp
            JOIN nodes_hierarchy nh_plan ON tp.id = nh_plan.id AND nh_plan.node_type_id = 5
            JOIN testprojects tpp ON tp.testproject_id = tpp.id
            JOIN nodes_hierarchy nh_proj ON tpp.id = nh_proj.id AND nh_proj.node_type_id = 1
            WHERE tp.id = %s
        """
        result = self.client._db_query(query, (plan_id,))
        return result[0] if result else {}
    
    def _get_build_info(self, build_id: int) -> Dict:
        query = "SELECT id, name, notes FROM builds WHERE id = %s"
        result = self.client._db_query(query, (build_id,))
        return result[0] if result else {}
    
    def _get_suite_structure(self, plan_id: int) -> List[Dict]:
        query = """
            SELECT DISTINCT nh.id, nh.name, ts.details
            FROM nodes_hierarchy nh
            JOIN testsuites ts ON nh.id = ts.id
            WHERE nh.node_type_id = 2
            AND nh.id IN (
                SELECT DISTINCT nh_suite.id
                FROM testplan_tcversions tptc
                JOIN tcversions tcv ON tptc.tcversion_id = tcv.id
                JOIN nodes_hierarchy nh_v ON tcv.id = nh_v.id AND nh_v.node_type_id = 4
                JOIN nodes_hierarchy nh_tc ON nh_v.parent_id = nh_tc.id AND nh_tc.node_type_id = 3
                JOIN nodes_hierarchy nh_suite ON nh_tc.parent_id = nh_suite.id AND nh_suite.node_type_id = 2
                WHERE tptc.testplan_id = %s
            )
        """
        return self.client._db_query(query, (plan_id,))
    
    def _format_status(self, status: str) -> str:
        status_map = {
            'p': 'PASS',
            'f': 'FAIL',
            'b': 'BLOCKED',
            'w': 'WARNING',
            'n': 'NOT RUN'
        }
        return status_map.get(status, status.upper() if status else 'NOT RUN')
    
    def _get_status_color(self, status: str) -> str:
        color_map = {
            'p': '70AD47',
            'f': 'C00000',
            'b': 'FF9900',
            'w': 'FFC000',
            'n': 'BFBFBF'
        }
        return color_map.get(status, 'BFBFBF')
    
    def _set_cell_background(self, cell, color: str):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    def _set_cell_text_color(self, cell, color: str):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor.from_string(color)
